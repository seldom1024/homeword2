from docx import Document
from os import path
from PIL import Image
import numpy as np
import jieba
from wordcloud import WordCloud, STOPWORDS


def handle_initialize(filename):
    doc = Document(filename)
    txt = ""
    for i in range(len(doc.paragraphs)):
        txt = txt+doc.paragraphs[i].text
    return " ".join(jieba.cut(txt))

filename = "D:\\作业\\大三\\数据挖掘\\实验三\\广东财经大学防控疫情确保开学安全工作方案.docx"
stopwords = set(STOPWORDS)
hh = handle_initialize(filename)
wc = WordCloud(background_color='white', max_font_size=100, font_path='./fonts/simhei.ttf',
               mask=np.array(Image.open(path.join(path.dirname(__file__), "1.png"))),
               stopwords=stopwords).generate(hh)
image = wc.to_image()
image.show()
